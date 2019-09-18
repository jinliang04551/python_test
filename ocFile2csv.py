#!/usr/bin/python2.7
#*-----encoding=utf8-----*

from docx.enum.style import WD_STYLE_TYPE
from docx import *

import re

f = open('./ZTBCmd.m')    
content = f.read()
print(content)

items = content.split(';')
print items
print len(items)

rows = []

for item in items:
    dic = {}
    startIndex =  item.find('/*')
    endIndex = item.find('*/')
    print 'startIndex:',startIndex
    print 'endIndex:',endIndex
    itemName = item[startIndex+2:endIndex]
    print itemName

    startIndex =  item.find('const')
    endIndex = item.find('ZTBCmd')
    itemType = item[startIndex+5:endIndex]  
    print itemType

    startIndex =  item.find('=')
    cmd = item[startIndex+1:]  
    print cmd
    
    dic['row1'] = itemName
    dic['row2'] = itemType
    dic['row3'] = cmd

    print('==============================')
    print dic
    print('==============================')

    rows.append(dic)

document = Document()

#样式表
# styles = document.styles

# for s  in styles:
#     if s.type  == WD_STYLE_TYPE.TABLE:
#         document.add_paragraph(u'表格样式' + s.name)
#         table = document.add_table(3,3,style = s)
#         heading_cells = table.rows[0].cells 
#         heading_cells[0].text = 'first row content'
#         heading_cells[1].text = 'first row content'
#         heading_cells[2].text = 'first row content'
#         document.add_paragraph("\n")

# document.save('demo.docx')

#建表 表头
row_count = len(rows) + 1
table = document.add_table(row_count,3,style='Medium Grid 1 Accent 1')
heading_cells = table.rows[0].cells 
heading_cells[0].text = u'第一列内容'
heading_cells[1].text = u'第二列内容'
heading_cells[2].text = u'第三列内容'

#赋值
for index  in range(len(rows)):
    dic = rows[index]
    itemName = dic['row1']
    itemType = dic['row2']
    itemCmd  = dic['row3']

    print('==============================')
    print('==============================')
    print dic
    print('==============================')     
    print('==============================')

    currentCells = table.rows[index+1].cells 

    currentCells[0].text = itemName.decode('utf-8')
    currentCells[1].text = itemType
    currentCells[2].text = itemCmd


document.save("demo.docx")


